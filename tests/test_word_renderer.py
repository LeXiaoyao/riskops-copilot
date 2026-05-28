"""Tests for the Word business report renderer."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from riskops.engines.report.word_renderer import write_business_report_word

_M3_SUMMARY = Path(__file__).resolve().parents[1] / "outputs" / "m3" / "m3_summary.json"
_ROI_RESULTS = Path(__file__).resolve().parents[1] / "outputs" / "model_lab" / "roi_results.json"


def test_write_business_report_word_creates_file(tmp_path: Path) -> None:
    output = tmp_path / "report.docx"

    result = write_business_report_word(_M3_SUMMARY, output, _ROI_RESULTS)

    assert output.exists()
    assert result["output_path"] == str(output)
    assert result["paragraph_count"] > 0


def test_word_report_has_section_headings(tmp_path: Path) -> None:
    output = tmp_path / "report.docx"

    write_business_report_word(_M3_SUMMARY, output, _ROI_RESULTS)

    doc = Document(str(output))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "核心结论" in full_text
    assert "异常信号" in full_text
    assert "归因分析" in full_text
    assert "ROI" in full_text
    assert "边界声明" in full_text


def test_word_report_contains_boundary_disclaimer(tmp_path: Path) -> None:
    output = tmp_path / "report.docx"

    write_business_report_word(_M3_SUMMARY, output, _ROI_RESULTS)

    doc = Document(str(output))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "synthetic data" in full_text
    assert "no real" in full_text.lower() or "no real financial" in full_text


def test_cli_render_word_can_run(tmp_path: Path) -> None:
    from riskops.interfaces.cli import main

    output = tmp_path / "report.docx"
    ret = main(
        [
            "render-word",
            "--input",
            str(_M3_SUMMARY),
            "--roi-input",
            str(_ROI_RESULTS),
            "--output",
            str(output),
        ]
    )

    assert ret == 0
    assert output.exists()
