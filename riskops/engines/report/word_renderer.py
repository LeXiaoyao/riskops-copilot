"""Word document renderer for the M4 business report."""

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from riskops.engines.report.business_report import BusinessReportInputError, load_m3_summary
from riskops.engines.report.excel_renderer import load_roi_results

HEADING_COLOR = RGBColor(0x1F, 0x4E, 0x79)
BODY_COLOR = RGBColor(0x1F, 0x1F, 0x1F)
CAPTION_COLOR = RGBColor(0x59, 0x59, 0x59)


class WordReportInputError(BusinessReportInputError):
    """Raised when Word report inputs are missing or malformed."""


def write_business_report_word(
    m3_summary_path: Path,
    output_path: Path,
    roi_results_path: Path,
) -> dict[str, Any]:
    """Generate a Word business report from M3 summary and ROI results.

    Returns metadata dict with output_path and paragraph_count.
    """
    summary = load_m3_summary(m3_summary_path)
    try:
        roi_results = load_roi_results(roi_results_path)
    except BusinessReportInputError as exc:
        raise WordReportInputError(str(exc)) from exc

    doc = Document()
    _set_default_font(doc)

    _add_title_section(doc)
    _add_executive_summary(doc, summary)
    _add_anomaly_section(doc, summary)
    _add_attribution_section(doc, summary)
    _add_roi_section(doc, roi_results)
    _add_boundary_section(doc)
    _add_next_actions_section(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    paragraph_count = len(doc.paragraphs)
    return {"output_path": str(output_path), "paragraph_count": paragraph_count}


# ---------------------------------------------------------------------------
# Sections
# ---------------------------------------------------------------------------


def _add_title_section(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RiskOps Copilot — 贷后经营复盘报告")
    run.bold = True
    run.font.size = Pt(20)
    _set_run_color(run, HEADING_COLOR)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"生成时间：{_generated_at_text()}")
    sub_run.font.size = Pt(10)
    _set_run_color(sub_run, CAPTION_COLOR)

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = caption.add_run(
        "⚠ 本报告基于 synthetic data（合成数据），仅用于 demo 演示，不代表真实业务结论。"
    )
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    _set_run_color(cap_run, CAPTION_COLOR)

    doc.add_paragraph()


def _add_executive_summary(doc: Document, summary: dict[str, Any]) -> None:
    overview = _as_dict(summary.get("anomaly_overview"))
    attribution = _as_dict(summary.get("m1_d7_attribution_summary"))
    target = _as_dict(attribution.get("target_anomaly")) or _as_dict(summary.get("attribution_target_anomaly"))
    severity = _as_dict(overview.get("severity_counts"))

    _heading(doc, "一、核心结论")
    lines = [
        f"M1 D7 回收率变动幅度：{_format_pct(target.get('relative_change'), signed=True)}",
        "异常数量：{} 条（high {} / medium {} / low {}）".format(
            _safe_int(overview.get("anomaly_count")),
            _safe_int(severity.get("high")),
            _safe_int(severity.get("medium")),
            _safe_int(severity.get("low")),
        ),
        f"归因目标指标：{attribution.get('target_metric_name_cn') or 'D7 回收率'}",
        "数据边界：synthetic data，不代表真实业务决策。",
    ]
    for line in lines:
        _bullet(doc, line)
    doc.add_paragraph()


def _add_anomaly_section(doc: Document, summary: dict[str, Any]) -> None:
    anomalies = [a for a in _as_list(summary.get("high_priority_anomalies")) if isinstance(a, dict)]

    _heading(doc, "二、高优先级异常信号")
    if not anomalies:
        _body(doc, "未发现高优先级异常信号。")
        doc.add_paragraph()
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(["指标", "严重度", "变动幅度", "建议动作"]):
        hdr[i].text = text
        _bold_cell(hdr[i])

    for anomaly in anomalies:
        row = table.add_row().cells
        row[0].text = str(anomaly.get("metric_name_cn") or "-")
        row[1].text = str(anomaly.get("severity") or "-")
        row[2].text = _format_pct(anomaly.get("relative_change"), signed=True)
        row[3].text = str(anomaly.get("recommended_next_step") or "-")

    doc.add_paragraph()


def _add_attribution_section(doc: Document, summary: dict[str, Any]) -> None:
    attribution = _as_dict(summary.get("m1_d7_attribution_summary"))
    drivers = [d for d in _as_list(attribution.get("top_drivers")) if isinstance(d, dict)][:5]

    _heading(doc, "三、归因分析（M1 D7 回收率）")
    if not drivers:
        _body(doc, "未发现归因因子。")
        doc.add_paragraph()
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(["维度", "值", "贡献度", "业务解读"]):
        hdr[i].text = text
        _bold_cell(hdr[i])

    for driver in drivers:
        row = table.add_row().cells
        row[0].text = str(driver.get("dimension_name") or "-")
        row[1].text = str(driver.get("dimension_value") or "-")
        row[2].text = _format_pct(driver.get("contribution_score"))
        row[3].text = str(driver.get("business_interpretation") or "-")

    doc.add_paragraph()
    _caption(
        doc,
        "归因边界：channel_code / province / score_band 为线索，不是最终根因；line_id 为作业队列，不是电话线路。",
    )
    doc.add_paragraph()


def _add_roi_section(doc: Document, roi_results: dict[str, Any]) -> None:
    rows = [r for r in _as_list(roi_results.get("results")) if isinstance(r, dict)]

    _heading(doc, "四、策略 ROI 测算")
    if not rows:
        _body(doc, "未发现 ROI 情景。")
        doc.add_paragraph()
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(["情景", "净收益", "ROI ratio", "回收期"]):
        hdr[i].text = text
        _bold_cell(hdr[i])

    for result in rows[:6]:
        row = table.add_row().cells
        row[0].text = str(result.get("scenario_name") or result.get("scenario_id") or "-")
        row[1].text = _format_money(result.get("net_benefit"))
        row[2].text = _format_number(result.get("roi_ratio"))
        row[3].text = _format_payback(result.get("payback_period_days"))

    doc.add_paragraph()
    _caption(doc, "ROI 边界：demo cost assumptions，不代表真实财务结果。")
    doc.add_paragraph()


def _add_boundary_section(doc: Document) -> None:
    _heading(doc, "五、AI + ML 边界声明")
    for line in [
        "synthetic data only — 本报告基于合成数据，不包含真实客户信息。",
        "no real financial conclusion — ROI 与收益数字来自 demo 假设。",
        "no real collection action — 不产生真实催收动作。",
        "no LLM automatic decisioning — LLM 输出仅供参考，需人工确认。",
        "no production claim — 不宣称真实上线或生产催收结果。",
    ]:
        _bullet(doc, line)
    doc.add_paragraph()


def _add_next_actions_section(doc: Document) -> None:
    _heading(doc, "六、建议后续动作")
    for line in [
        "人工复核 M1 D7 回收率口径与观察窗口期。",
        "下钻 Top5 归因因子的渠道、区域、客群和过程证据。",
        "复核 ROI 情景的成本假设与收益假设，确认商务边界。",
        "策略动作进入人工评审流程，不触发真实催收执行。",
        "补充业务 owner 结论和数据口径确认后，再进入对外材料。",
    ]:
        _bullet(doc, line)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)


def _heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    _set_run_color(run, HEADING_COLOR)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)


def _body(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(4)


def _bullet(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Bullet")
    para.text = text
    para.paragraph_format.space_after = Pt(2)


def _caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    _set_run_color(run, CAPTION_COLOR)


def _bold_cell(cell) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


def _set_run_color(run, color: RGBColor) -> None:
    run.font.color.rgb = color


def _generated_at_text() -> str:
    return datetime.now(UTC).replace(microsecond=0).isoformat()


def _format_pct(value: Any, *, signed: bool = False) -> str:
    if not isinstance(value, int | float):
        return "-"
    pct = value * 100
    prefix = "+" if signed and pct > 0 else ""
    return f"{prefix}{pct:.2f}%"


def _format_number(value: Any) -> str:
    if not isinstance(value, int | float):
        return "-"
    return f"{value:.2f}"


def _format_money(value: Any) -> str:
    if not isinstance(value, int | float):
        return "-"
    return f"{value:,.2f}"


def _format_payback(value: Any) -> str:
    if not isinstance(value, int | float):
        return "-"
    if value < 1:
        return "<1 天"
    return f"{value:.1f} 天"


def _as_dict(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _as_list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []


def _safe_int(value: Any) -> int:
    if isinstance(value, bool):
        return int(value)
    if isinstance(value, int | float):
        return int(value)
    return 0
